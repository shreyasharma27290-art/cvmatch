"""
CVMatch Resume Parser
Extracts structured data from PDF and DOCX resume files.

Dependencies:
  pip install pdfplumber python-docx spacy --break-system-packages
  python -m spacy download en_core_web_sm
"""

import re
import io
import logging
from typing import Dict, List, Optional, Tuple
from pathlib import Path

logger = logging.getLogger(__name__)

# ---------------------------------------------------------------------------
# Known skill dictionary (extend as needed)
# ---------------------------------------------------------------------------

SKILLS_DB = {
    "languages": [
        "Python", "Java", "JavaScript", "TypeScript", "C++", "C#", "C", "Go",
        "Rust", "Kotlin", "Swift", "R", "MATLAB", "Scala", "PHP", "Ruby",
    ],
    "web_frameworks": [
        "React", "React.js", "Angular", "Vue.js", "Next.js", "Svelte",
        "Node.js", "Express.js", "Django", "FastAPI", "Flask", "Spring Boot",
        "Laravel", "Ruby on Rails", "ASP.NET", "Nuxt.js",
    ],
    "ml_ai": [
        "Machine Learning", "Deep Learning", "NLP", "Computer Vision",
        "TensorFlow", "PyTorch", "Keras", "Scikit-learn", "OpenCV",
        "HuggingFace", "LangChain", "LLM", "RAG", "YOLO",
        "Pandas", "NumPy", "Matplotlib", "Seaborn", "Plotly",
        "Apache Spark", "Kafka", "Airflow", "dbt",
    ],
    "databases": [
        "SQL", "PostgreSQL", "MySQL", "MongoDB", "Redis", "SQLite",
        "Cassandra", "Elasticsearch", "DynamoDB", "Firebase",
        "Oracle", "Microsoft SQL Server",
    ],
    "cloud_devops": [
        "AWS", "GCP", "Azure", "Docker", "Kubernetes", "Terraform",
        "CI/CD", "Jenkins", "GitHub Actions", "ArgoCD", "Helm",
        "Linux", "Bash", "Shell Scripting", "Nginx", "Ansible",
    ],
    "tools": [
        "Git", "GitHub", "GitLab", "Bitbucket", "Postman", "Jira",
        "Figma", "VS Code", "IntelliJ", "Jupyter Notebook",
        "Power BI", "Tableau", "Looker", "Excel",
    ],
    "soft_skills": [
        "Problem Solving", "Team Collaboration", "Communication",
        "Leadership", "Project Management", "Agile", "Scrum",
    ],
}

ALL_SKILLS: List[str] = [s for lst in SKILLS_DB.values() for s in lst]

# Section header patterns
SECTION_PATTERNS = {
    "contact":    re.compile(r"(contact|phone|email|linkedin|github|portfolio)", re.I),
    "summary":    re.compile(r"(summary|objective|profile|about me|career goal)", re.I),
    "education":  re.compile(r"(education|qualification|academic|degree|university|college|school)", re.I),
    "experience": re.compile(r"(experience|work history|employment|internship|career)", re.I),
    "skills":     re.compile(r"(skill|technical|competenc|proficien|technology|stack)", re.I),
    "projects":   re.compile(r"(project|portfolio|side project|personal project|open.?source)", re.I),
    "certifications": re.compile(r"(certif|course|training|award|achievement|accomplishment)", re.I),
    "languages":  re.compile(r"(language|spoken|fluent)", re.I),
}


# ---------------------------------------------------------------------------
# Extraction helpers
# ---------------------------------------------------------------------------

def extract_email(text: str) -> Optional[str]:
    m = re.search(r'\b[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Za-z]{2,}\b', text)
    return m.group(0) if m else None


def extract_phone(text: str) -> Optional[str]:
    m = re.search(r'(\+91[-\s]?)?[6-9]\d{9}', text)
    return m.group(0) if m else None


def extract_linkedin(text: str) -> Optional[str]:
    m = re.search(r'linkedin\.com/in/[\w-]+', text, re.I)
    return f"https://www.{m.group(0)}" if m else None


def extract_github(text: str) -> Optional[str]:
    m = re.search(r'github\.com/[\w-]+', text, re.I)
    return f"https://www.{m.group(0)}" if m else None


def extract_name(lines: List[str]) -> str:
    """Heuristic: the resume's first non-empty line is typically the name."""
    for line in lines:
        line = line.strip()
        if line and len(line.split()) <= 5 and not re.search(r'[@|/\\]', line):
            # Looks like a name (short, no special chars)
            if re.match(r'^[A-Za-z\s.]+$', line):
                return line
    return "Unknown"


def extract_skills(text: str) -> Dict[str, List[str]]:
    """Return skills grouped by category."""
    text_lower = text.lower()
    result: Dict[str, List[str]] = {}

    for category, skills in SKILLS_DB.items():
        found = []
        for skill in skills:
            pattern = r'\b' + re.escape(skill.lower()) + r'\b'
            if re.search(pattern, text_lower):
                found.append(skill)
        if found:
            result[category] = found

    return result


def extract_experience_years(text: str) -> int:
    """Estimate total years of experience from the resume text."""
    # Look for patterns like "3 years", "2+ years", etc.
    patterns = [
        r'(\d+)\+?\s+years?\s+of\s+experience',
        r'(\d+)\+?\s+yrs?\s+of\s+experience',
        r'experience\s+of\s+(\d+)\+?\s+years?',
    ]
    for p in patterns:
        m = re.search(p, text, re.I)
        if m:
            return int(m.group(1))

    # Count internship/work entries as rough proxy
    internship_count = len(re.findall(r'intern', text, re.I))
    job_count = len(re.findall(r'\b(engineer|analyst|developer|manager|lead)\b', text, re.I))

    if internship_count >= 2 or job_count >= 3:
        return 1
    return 0


def extract_education(text: str) -> List[Dict]:
    """Extract degree, institution, year from education section."""
    educations = []
    degree_patterns = [
        r'(b\.?tech|bachelor|b\.?e\.?|b\.?sc|b\.?com)',
        r'(m\.?tech|master|m\.?sc|mba|m\.?e\.?)',
        r'(ph\.?d|doctorate)',
        r'(diploma|polytechnic)',
    ]

    lines = text.split('\n')
    for i, line in enumerate(lines):
        line_lower = line.lower()
        for pattern in degree_patterns:
            if re.search(pattern, line_lower, re.I):
                # Try to extract year
                year_match = re.search(r'(19|20)\d{2}', line)
                cgpa_match = re.search(r'(cgpa|gpa|percentage|%)[:\s]*([\d.]+)', line, re.I)

                educations.append({
                    "degree": line.strip()[:120],
                    "year": year_match.group(0) if year_match else None,
                    "cgpa": cgpa_match.group(2) if cgpa_match else None,
                })
                break

    return educations[:4]  # cap at 4


def extract_projects(text: str) -> List[str]:
    """Extract project names/descriptions."""
    projects = []
    lines = text.split('\n')
    in_projects = False

    for line in lines:
        line_stripped = line.strip()
        if not line_stripped:
            continue

        if SECTION_PATTERNS["projects"].search(line_stripped) and len(line_stripped) < 30:
            in_projects = True
            continue

        # Stop at the next major section
        if in_projects and any(
            p.search(line_stripped) and len(line_stripped) < 30
            for key, p in SECTION_PATTERNS.items()
            if key != "projects"
        ):
            break

        if in_projects and len(line_stripped) > 10:
            projects.append(line_stripped[:200])

    return projects[:6]


def detect_sections(text: str) -> Dict[str, bool]:
    """Check which standard sections are present."""
    text_lower = text.lower()
    return {
        key: bool(pattern.search(text_lower))
        for key, pattern in SECTION_PATTERNS.items()
    }


# ---------------------------------------------------------------------------
# PDF & DOCX readers
# ---------------------------------------------------------------------------

def read_pdf(file_bytes: bytes) -> str:
    """Extract text from PDF using pdfplumber (best layout fidelity)."""
    try:
        import pdfplumber
        with pdfplumber.open(io.BytesIO(file_bytes)) as pdf:
            pages = [page.extract_text() or "" for page in pdf.pages]
        return "\n".join(pages)
    except ImportError:
        logger.warning("pdfplumber not installed. Falling back to PyPDF2.")
        return _read_pdf_pypdf2(file_bytes)
    except Exception as e:
        logger.error(f"PDF read error: {e}")
        return ""


def _read_pdf_pypdf2(file_bytes: bytes) -> str:
    try:
        import PyPDF2
        reader = PyPDF2.PdfReader(io.BytesIO(file_bytes))
        return "\n".join(page.extract_text() or "" for page in reader.pages)
    except Exception as e:
        logger.error(f"PyPDF2 error: {e}")
        return ""


def read_docx(file_bytes: bytes) -> str:
    """Extract text from DOCX using python-docx."""
    try:
        from docx import Document
        doc = Document(io.BytesIO(file_bytes))
        paragraphs = [p.text for p in doc.paragraphs]
        # Also extract from tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    paragraphs.append(cell.text)
        return "\n".join(paragraphs)
    except ImportError:
        logger.warning("python-docx not installed.")
        return ""
    except Exception as e:
        logger.error(f"DOCX read error: {e}")
        return ""


# ---------------------------------------------------------------------------
# Master parse function
# ---------------------------------------------------------------------------

def parse_resume(file_bytes: bytes, file_type: str) -> Dict:
    """
    Main entry-point. Returns a structured dict with all extracted info.

    Args:
        file_bytes: raw bytes of the uploaded file
        file_type:  "pdf" | "docx" | "txt"

    Returns:
        {
          "raw_text": str,
          "contact": {...},
          "skills": {...categorised...},
          "all_skills": [...flat list...],
          "education": [...],
          "projects": [...],
          "experience_years": int,
          "sections_found": {...},
          "word_count": int,
          "page_estimate": int,
        }
    """
    # Step 1 – Extract raw text
    if file_type == "pdf":
        raw_text = read_pdf(file_bytes)
    elif file_type in ("docx", "doc"):
        raw_text = read_docx(file_bytes)
    elif file_type == "txt":
        raw_text = file_bytes.decode("utf-8", errors="ignore")
    else:
        raw_text = file_bytes.decode("utf-8", errors="ignore")

    if not raw_text.strip():
        logger.warning("Empty text extracted from resume — possibly image-based PDF")
        raw_text = "[Could not extract text — please use a text-based PDF or DOCX]"

    lines = [l for l in raw_text.split('\n') if l.strip()]

    # Step 2 – Extract fields
    skills_by_category = extract_skills(raw_text)
    all_skills_flat: List[str] = [s for lst in skills_by_category.values() for s in lst]

    result = {
        "raw_text": raw_text,
        "word_count": len(raw_text.split()),
        "page_estimate": max(1, len(raw_text) // 2500),
        "contact": {
            "name": extract_name(lines),
            "email": extract_email(raw_text),
            "phone": extract_phone(raw_text),
            "linkedin": extract_linkedin(raw_text),
            "github": extract_github(raw_text),
        },
        "skills": skills_by_category,
        "all_skills": all_skills_flat,
        "education": extract_education(raw_text),
        "projects": extract_projects(raw_text),
        "experience_years": extract_experience_years(raw_text),
        "sections_found": detect_sections(raw_text),
    }

    logger.info(
        f"Parsed resume: {result['contact']['name']} | "
        f"{len(all_skills_flat)} skills | "
        f"~{result['page_estimate']} page(s)"
    )
    return result


# ---------------------------------------------------------------------------
# Quick self-test
# ---------------------------------------------------------------------------

if __name__ == "__main__":
    SAMPLE = """
Priya Sharma
priya@example.com | +91 9876543210 | linkedin.com/in/priya | github.com/priya-dev

EDUCATION
B.Tech Information Technology | Lovely Professional University | 2021-2025 | CGPA: 8.2

SKILLS
Python, JavaScript, React.js, FastAPI, TensorFlow, SQL, PostgreSQL, Docker, AWS, Git

EXPERIENCE
Data Science Intern | ABC Analytics | June 2024 - Aug 2024
- Built ML models with Python and Scikit-learn, improving accuracy by 18%
- Analysed 100K+ records using Pandas and SQL

PROJECTS
AI Recommendation Engine
- Collaborative filtering system using TensorFlow, served 50K+ users
- Deployed on AWS with 99.9% uptime

CERTIFICATIONS
AWS Cloud Practitioner | 2024
Python for Data Science – Coursera | 2024
"""

    parsed = parse_resume(SAMPLE.encode(), file_type="txt")
    print("\n=== PARSED RESUME ===")
    print(f"Name      : {parsed['contact']['name']}")
    print(f"Email     : {parsed['contact']['email']}")
    print(f"Skills    : {parsed['all_skills']}")
    print(f"Edu       : {parsed['education']}")
    print(f"Exp yrs   : {parsed['experience_years']}")
    print(f"Projects  : {parsed['projects'][:2]}")
    print(f"Sections  : {parsed['sections_found']}")
